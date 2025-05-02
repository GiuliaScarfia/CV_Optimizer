import os
from docx import Document
import logging
from io import BytesIO


# Funzione per modificare il template Word
def modify_template(dati_cv):
    num_experience = len(dati_cv["esperienze"])

    base_dir = os.path.dirname(os.path.abspath(__file__))  # path di utils/
    data_dir = os.path.abspath(os.path.join(base_dir, '..', 'data'))  # path assoluto della cartella data/

    if num_experience == 1:
        template_path = os.path.join(data_dir, 'template_curriculum_1.docx')
    elif num_experience == 2:
        template_path = os.path.join(data_dir, 'template_curriculum_2.docx')
    elif num_experience >= 3:
        template_path = os.path.join(data_dir, 'template_curriculum_3.docx')
    else:
        logging.error("Errore: Numero di esperienze non supportato.")
        exit(1)

    doc = Document(template_path)
    return doc

    # Dizionario dei segnaposto e dei valori da sostituire
    label_values = {
        "{nome}": dati_cv["nome"],
        "{cognome}": dati_cv["cognome"],
        "{posizione}": dati_cv["posizione"],
        "{telefono}": dati_cv["telefono"],
        "{email}": dati_cv["email"],
        "{sito}": dati_cv["sito"],
        "{città}": dati_cv["città"],
        "{obiettivo}": dati_cv["obiettivo"],
        "{competenze}": "\n".join([f"• {competenza.strip()}" for competenza in dati_cv["competenze"].split(",")]),
        "{lingue}": "\n".join([f"• {lingua.strip()}" for lingua in dati_cv["lingue"].split(",")]),
        "{comunicazione}": dati_cv["comunicazione"],
        "{hobby}": dati_cv["hobby"]
    }

    # Aggiungiamo i segnaposto per i titoli di studio
    for i in range(1, 3):  # Supponendo sempre 2 blocchi per istruzione
        if 0 <= i-1 < len(dati_cv["istruzione"]):
            education = dati_cv["istruzione"][i-1]
        else:
            education = {}  # or some default value

        label_values[f"{{data_conseguimento_{i}}}"] = education["data_conseguimento"]
        label_values[f"{{nome_istituto_{i}}}"] = education["nome_istituto"]
        label_values[f"{{titolo_{i}}}"] = education["titolo"]
        label_values[f"{{desc_titolo_{i}}}"] = education["desc_titolo"]

    # Aggiungi i segnaposto per le esperienze
    for i in range(1, 4):  # Supponendo 3 blocchi per le esperienze
        if i <= len(dati_cv["esperienze"]):
            experience = dati_cv["esperienze"][i - 1]
            label_values[f"{{periodo_{i}}}"] = experience["periodo"]
            label_values[f"{{qualifica_{i}}}"] = experience["qualifica"]
            label_values[f"{{nome_società_{i}}}"] = experience["nome_società"]
            label_values[f"{{desc_esperienza_{i}}}"] = experience["desc_esperienza"]
        else:
            label_values[f"{{periodo_{i}}}"] = ""
            label_values[f"{{qualifica_{i}}}"] = ""
            label_values[f"{{nome_società_{i}}}"] = ""
            label_values[f"{{desc_esperienza_{i}}}"] = ""

    # Funzione per sostituire i segnaposto in un paragrafo mantenendo la formattazione
    def replace_label_with_format(paragraph, label_values):
        # Unisci il testo di tutti i runs
        full_text = "".join(run.text for run in paragraph.runs)

        # Effettua le sostituzioni sul testo completo
        for label, value in label_values.items():
            if label in full_text:  # Verifica se il segnaposto è presente
                full_text = full_text.replace(label, value)

        # Aggiorna i runs esistenti con il nuovo testo
        if paragraph.runs:
            paragraph.runs[0].text = full_text  # Aggiorna il primo run con il testo completo
            for run in paragraph.runs[1:]:
                run.text = ""  # Svuota gli altri runs

    # Sostituisci i segnaposto nei paragrafi
    for paragraph in doc.paragraphs:
        replace_label_with_format(paragraph, label_values)

    # Sostituisci i segnaposto nelle tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_label_with_format(paragraph, label_values)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()  # Restituisce il contenuto del file come bytes



